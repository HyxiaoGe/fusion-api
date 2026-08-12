from __future__ import annotations

import csv
import io
import multiprocessing
import unicodedata
import zipfile
from contextlib import suppress
from dataclasses import dataclass
from pathlib import PurePath


class KnowledgeParseError(ValueError):
    def __init__(self, code: str, summary: str):
        self.code = code
        self.summary = summary
        super().__init__(summary)


@dataclass(frozen=True)
class ParsedSection:
    text: str
    page: int | None = None
    section: str | None = None


def _isolated_parse_child(connection, content: bytes, mimetype: str, filename: str) -> None:
    try:
        import resource

        memory_limit = 512 * 1024 * 1024
        with suppress(Exception):
            resource.setrlimit(resource.RLIMIT_AS, (memory_limit, memory_limit))
        with suppress(Exception):
            resource.setrlimit(resource.RLIMIT_CPU, (45, 45))
        connection.send(("ok", KnowledgeDocumentParser().parse(content, mimetype=mimetype, filename=filename)))
    except KnowledgeParseError as exc:
        connection.send(("error", (exc.code, exc.summary)))
    except Exception:
        connection.send(("error", ("KNOWLEDGE_DOCUMENT_INVALID", "文档解析子进程异常")))
    finally:
        connection.close()


def parse_document_isolated(
    content: bytes,
    *,
    mimetype: str,
    filename: str,
    timeout_seconds: int,
) -> list[ParsedSection]:
    """在有限资源子进程中解析复杂容器格式。"""
    context = multiprocessing.get_context("spawn")
    parent, child = context.Pipe(duplex=False)
    process = context.Process(
        target=_isolated_parse_child,
        args=(child, content, mimetype, filename),
        daemon=True,
    )
    process.start()
    child.close()
    try:
        if not parent.poll(timeout_seconds):
            process.terminate()
            process.join(timeout=5)
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_PARSE_TIMEOUT", "文档解析超过时间上限")
        status, payload = parent.recv()
        process.join(timeout=5)
        if status == "ok":
            return payload
        code, summary = payload
        raise KnowledgeParseError(code, summary)
    finally:
        parent.close()
        if process.is_alive():
            process.terminate()
            process.join(timeout=5)


class KnowledgeDocumentParser:
    """v1 确定性文本解析器；不调用 LLM、视觉模型或 OCR。"""

    VERSION = "parser-v1"
    MAX_CHARACTERS = 2_000_000
    MAX_PAGES = 1000
    MAX_DOCX_ENTRIES = 5000
    MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
    MIME_SUFFIXES = {
        "text/plain": {".txt", ".text"},
        "text/markdown": {".md", ".markdown"},
        "text/csv": {".csv"},
        "application/pdf": {".pdf"},
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {".docx"},
    }

    def parse(self, content: bytes, *, mimetype: str, filename: str) -> list[ParsedSection]:
        suffix = PurePath(filename).suffix.lower()
        allowed_suffixes = self.MIME_SUFFIXES.get(mimetype)
        if allowed_suffixes is None:
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_UNSUPPORTED", "当前知识库版本不支持该文档格式")
        if suffix not in allowed_suffixes:
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_TYPE_MISMATCH", "文档 MIME 与扩展名不一致")
        if mimetype == "application/pdf" and not content.startswith(b"%PDF-"):
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_TYPE_MISMATCH", "PDF 文件签名无效")
        if mimetype.endswith("wordprocessingml.document") and not content.startswith(b"PK"):
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_TYPE_MISMATCH", "DOCX 文件签名无效")
        if mimetype in {"text/plain", "text/markdown"}:
            sections = [ParsedSection(self._decode_text(content))]
        elif mimetype == "text/csv":
            sections = [ParsedSection(self._parse_csv(content))]
        elif mimetype == "application/pdf":
            sections = self._parse_pdf(content)
        elif mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            sections = self._parse_docx(content)
        else:
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_UNSUPPORTED", "当前知识库版本不支持该文档格式")
        normalized = []
        for section in sections:
            if "\x00" in section.text:
                raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_INVALID", "文档正文包含不支持的 NUL 字符")
            text = self._normalize_text(section.text)
            if text:
                normalized.append(ParsedSection(text, section.page, section.section))
        total_characters = sum(len(section.text) for section in normalized)
        if not normalized or total_characters == 0:
            code = (
                "KNOWLEDGE_SCANNED_DOCUMENT_UNSUPPORTED"
                if mimetype == "application/pdf"
                else "KNOWLEDGE_DOCUMENT_EMPTY"
            )
            summary = (
                "PDF 未检测到文字层，知识库 v1 不支持 OCR" if mimetype == "application/pdf" else "文档未包含可索引文本"
            )
            raise KnowledgeParseError(code, summary)
        if total_characters > self.MAX_CHARACTERS:
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_TOO_LARGE", "文档解压后的文本内容超过处理上限")
        return normalized

    @staticmethod
    def _decode_text(content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_ENCODING_UNSUPPORTED", "文本文件必须使用 UTF-8 编码")

    def _parse_csv(self, content: bytes) -> str:
        decoded = self._decode_text(content)
        try:
            rows = csv.reader(io.StringIO(decoded, newline=""))
            return "\n".join("\t".join(cell.strip() for cell in row) for row in rows)
        except csv.Error as exc:
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_INVALID", "CSV 文档格式无效") from exc

    def _parse_pdf(self, content: bytes) -> list[ParsedSection]:
        import PyPDF2

        try:
            reader = PyPDF2.PdfReader(io.BytesIO(content), strict=True)
            if reader.is_encrypted:
                raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_ENCRYPTED", "不支持加密 PDF")
            if len(reader.pages) > self.MAX_PAGES:
                raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_TOO_LARGE", "PDF 页数超过处理上限")
            return [ParsedSection(page.extract_text() or "", page=index) for index, page in enumerate(reader.pages, 1)]
        except KnowledgeParseError:
            raise
        except Exception as exc:
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_INVALID", "PDF 文档无法解析") from exc

    @classmethod
    def _parse_docx(cls, content: bytes) -> list[ParsedSection]:
        import docx

        try:
            cls._validate_docx_archive(content)
            document = docx.Document(io.BytesIO(content))
            sections = [ParsedSection(paragraph.text, section="paragraph") for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    sections.append(ParsedSection("\t".join(cell.text for cell in row.cells), section="table"))
            return sections
        except KnowledgeParseError:
            raise
        except Exception as exc:
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_INVALID", "DOCX 文档无法解析") from exc

    @classmethod
    def _validate_docx_archive(cls, content: bytes) -> None:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > cls.MAX_DOCX_ENTRIES:
                    raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_TOO_LARGE", "DOCX 文件条目数超过处理上限")
                if any(entry.flag_bits & 0x1 for entry in entries):
                    raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_ENCRYPTED", "不支持加密 DOCX")
                if sum(entry.file_size for entry in entries) > cls.MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_TOO_LARGE", "DOCX 解压大小超过处理上限")
        except KnowledgeParseError:
            raise
        except (zipfile.BadZipFile, OSError) as exc:
            raise KnowledgeParseError("KNOWLEDGE_DOCUMENT_INVALID", "DOCX 文档无法解析") from exc

    @staticmethod
    def _normalize_text(text: str) -> str:
        normalized = unicodedata.normalize("NFKC", text).replace("\r\n", "\n").replace("\r", "\n")
        lines = [" ".join(line.split()) for line in normalized.split("\n")]
        return "\n".join(lines).strip()
